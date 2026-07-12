from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "results" / "reports" / "PAC-Test_v2.1_问题回应与修改建议_20260712.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size: float, bold=False, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before: float, after: float, line: float = 1.1) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_text(doc: Document, text: str, *, bold=False, color: str | None = None) -> None:
    p = doc.add_paragraph(style="Normal")
    style_paragraph(p, 0, 6, 1.1)
    run = p.add_run(text)
    set_font(run, 11, bold, color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, 0, 5, 1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_font(run, 11)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_font(run, {1: 16, 2: 13, 3: 12}[level], True, BLUE if level < 3 else DARK_BLUE)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    style_paragraph(p, 0, 0, 1.15)
    run = p.add_run(label + " ")
    set_font(run, 11, True, DARK_BLUE)
    run = p.add_run(text)
    set_font(run, 11, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, 0, 0, 1.1)
        run = p.add_run(value)
        set_font(run, 10, True, DARK_BLUE)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, 0, 0, 1.1)
            run = p.add_run(value)
            set_font(run, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, 0, 0, 1.0)
    run = p.add_run("PAC-Test v2.1 | 问题回应与修改建议（审阅稿）")
    set_font(run, 9, False, "6B7280")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, 0, 4, 1.0)
    run = title.add_run("PAC-Test v2.1 问题回应与修改建议")
    set_font(run, 20, True, DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, 0, 14, 1.0)
    run = subtitle.add_run("基于当前已完成结果与原始返回内容的核查稿 | 2026-07-12")
    set_font(run, 10, False, "6B7280")

    add_callout(
        doc,
        "核查结论：",
        "字段保持与干扰捕获率均可由现有评分逻辑复现；Qwen3-8B 的异常低分不是 API 调用失败，而是大量输出未遵守“只输出最终代码”的格式要求，并在 192 token 上限处截断；PAC-D 热图的 0、0.5、1 来自每个条件仅 2 条样本的离散均值。建议保留原始结果，但对其解释和图表呈现作透明修订。",
    )

    add_heading(doc, "一、核查范围与结论边界", 1)
    add_text(doc, "本次核查使用当前 PAC v2.1 的去重汇总结果、条件级汇总表、错误类型表和 Qwen3-8B 原始返回内容。核查目的不是重跑或改写原始分数，而是确认指标计算、异常结果与图表呈现是否存在误读风险。")
    add_bullet(doc, "本回应不修改任何原始预测、得分或 API 返回记录。")
    add_bullet(doc, "建议修改的是论文中的方法说明、图表表达和结论边界，而不是用人工方式修正模型结果。")

    add_heading(doc, "二、字段保持与干扰捕获率的计算口径", 1)
    add_heading(doc, "1. 字段保持（Field Retention）", 2)
    add_text(doc, "每条样本都包含 m 个必须同时保持正确绑定关系的目标字段，例如三段代码或多实体对应值。对一条模型输出，先在完整输出文本中逐项检查是否命中了每一个目标值，再以命中字段数除以字段总数得到该样本的字段准确率：")
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(formula, 4, 6, 1.0)
    run = formula.add_run("Field Accuracy_i = 正确命中的目标字段数 / 该样本目标字段总数")
    set_font(run, 11, True, DARK_BLUE)
    add_text(doc, "表格中的“字段保持”不是只看完整答对的样本，而是先对每个子实验内的字段准确率取平均，再对 PAC-A、PAC-B、PAC-C、PAC-D 四个子实验做等权平均并转换为百分比。它反映模型在完整答案失败时，是否仍保留了部分实体-属性或链路字段的正确绑定。评分时只使用 API 正常返回的样本。")

    add_heading(doc, "2. 干扰捕获率（Decoy Capture Rate）", 2)
    add_text(doc, "每条样本在正确证据旁放入了具有高相似形式、但实体、状态、批次、时间、通道或验证条件不满足要求的干扰值。若模型输出未构成完整正确答案，却包含任一预设干扰值，该样本被标记为 decoy_value_capture。子实验层面的干扰捕获率为：")
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(formula, 4, 6, 1.0)
    run = formula.add_run("Decoy Capture Rate_s = 干扰值被错误选入的样本数 / 子实验 s 的正常返回样本数")
    set_font(run, 11, True, DARK_BLUE)
    add_text(doc, "表格中的“干扰捕获率”同样是四个子实验率的等权平均，数值越低越好。该指标与字段保持互补：前者描述模型是否被相似但不合格的证据带偏，后者描述模型保留了多少正确字段。论文中应在首次出现该表时补充以上两句定义，并注明两个汇总指标均为子实验等权平均。")

    add_heading(doc, "三、Qwen3-8B 全部或接近全部为 0 的原因", 1)
    add_text(doc, "结论是：这不是 API 未返回导致的 0，也不能直接当成 Qwen3-8B 完全没有长上下文推理能力。它是当前“精确答案 + 192 token 输出上限”评测配置下，输出格式失配与截断共同造成的异常表现。")
    add_table(doc, ["核查项", "核查结果", "含义"], [
        ["调用是否成功", "231 / 231 条均为正常 API 返回；API error 为 0", "排除“调用失败后按 0 计分”的可能。"],
        ["精确正确数", "PAC-A 0 / 75；PAC-B 3 / 72；PAC-C 0 / 72；PAC-D 0 / 12", "完整答案严格匹配时只有 PAC-B 的 3 条成功。"],
        ["输出长度触顶", "200 / 231 条（86.6%）达到 192 token 上限", "大量回答在给出最终代码前被截断。"],
        ["原始输出形态", "多数样本先输出英文分步解释、证据复述和推理过程，而非直接输出 V1|V2|V3", "与任务要求的“只输出最终答案”不一致，降低了精确匹配得分。"],
    ], [1750, 3600, 4010])
    add_text(doc, "评分器已在完整返回文本中检索目标代码，并非只取最后一行；因此这不是简单的“答案在前面却没有被识别”。但在输出被截断的样本中，模型通常没有给出完整三段答案，或未能保留足够的正确字段。当前结果可被保留为“该解码配置下的端到端任务表现”，但不宜单独用作模型内在长上下文能力的强结论。")
    add_callout(doc, "建议的论文处理：", "保留 Qwen3-8B 原始数据与附表，不建议悄悄删改；在正文主比较图中将其标注为“输出格式/长度受限的诊断样本”，或移至附录并明确说明原因。若将来有预算，只对该模型用更严格的结构化输出约束和更高输出上限复测，才能作公平的能力比较。")

    add_heading(doc, "四、PAC-C / PAC-D 热图为何出现 0、0.5、1", 1)
    add_text(doc, "这不是热图程序错误，而是条件单元样本量较小造成的离散均值。PAC-C 每个“绑定容量 K / 查询数 Q”条件下有 8 条样本，因此准确率以 1/8 = 0.125 为最小步长；PAC-D 每个“跳数 / 假链数”条件下只有 2 条样本，因此准确率只能以 1/2 = 0.5 为最小步长。")
    add_table(doc, ["图表", "条件级重复数", "数值表现", "可否作为主结论"], [
        ["PAC-C 绑定容量热图", "通常为 8 条 / 单元", "可出现 0、0.125、0.25、…、1", "可用于展示趋势，但应避免过度解释单个格子的细微差别。"],
        ["PAC-D 多跳假链热图", "通常为 2 条 / 单元；Qwen3-14B-T 个别单元更少", "只能出现 0、0.5、1", "应定位为探索性或有效性验证，不宜作为精细的模型排序证据。"],
    ], [2250, 1900, 2100, 3110])
    add_text(doc, "因此，图 4-6 中大量 0、0.5、1 的根源是 PAC-D 的设计规模，而不是计算错误。论文修改时建议将 PAC-D 热图替换为“按模型汇总的平均准确率 + 条件说明”的柱状图或点图，并在图注中写明“每个条件重复 2 次”。这样既保持结果透明，也避免视觉上像连续精确估计。")

    add_heading(doc, "五、图表重叠与后续修改建议", 1)
    add_text(doc, "当前“PAC v2.1 error tendency”图中，图例位于下方并与横轴模型标签/横轴标题发生视觉重叠。这是排版问题，不涉及数值计算。Qwen3-8B 在该图中仍会显得异常，原因与上一节相同：大量格式失配和输出截断导致其 exact accuracy 极低。")
    add_table(doc, ["问题", "建议修改", "目的"], [
        ["图例与横轴文字重叠", "把图例移到图框外下方，增加底部边距，或改为图上直接标注。", "保证图可直接插入论文且不遮挡标签。"],
        ["PAC-D 离散热图", "正文改为模型级汇总图；条件级热图保留在附录，并在图注注明每格 2 条样本。", "避免把小样本离散均值误读为稳定曲线。"],
        ["Qwen3-8B 异常低分", "正文单列“输出格式失配诊断”，原始结果与说明放附录。", "透明呈现异常，同时不把协议问题误写成纯能力差距。"],
        ["指标定义缺失", "在 PAC-Test 方法或表注中补充字段保持、干扰捕获率、精确准确率的定义和方向。", "使读者能够自行理解表中的综合指标。"],
    ], [2000, 4200, 3160])

    add_heading(doc, "六、可直接补入论文的说明文字", 1)
    add_heading(doc, "1. 指标说明（表格前或表注）", 2)
    quote = "除完整答案准确率外，本文报告字段保持率与干扰捕获率。字段保持率计算模型输出中被正确保留的目标字段比例，用于刻画完整答案失败时的部分绑定保持能力；干扰捕获率计算模型将预设高相似干扰值误选入输出的比例，数值越低代表抗干扰能力越强。两项指标均先在各子实验内计算，再在 PAC-A 至 PAC-D 之间进行等权汇总。"
    add_callout(doc, "建议文字：", quote)
    add_heading(doc, "2. PAC-D 图注或正文说明", 2)
    quote = "PAC-D 用于验证多跳假链条件下的链路追踪失效。由于本轮每个跳数与假链数条件仅包含 2 条测试样本，条件级结果呈现为 0、0.5 和 1 等离散取值；因此本文将其作为机制性探索证据，并以模型级汇总结果作为主要描述，不对单个条件格作过度统计推断。"
    add_callout(doc, "建议文字：", quote)
    add_heading(doc, "3. Qwen3-8B 结果说明", 2)
    quote = "Qwen3-8B 在当前协议下表现出明显的输出格式失配：其多数返回以长推理文本展开，并频繁达到输出长度上限，导致完整代码串未能按要求生成。该结果反映该模型在本评测协议下的端到端可用性限制，本文将其与其他模型的任务能力比较分开讨论。"
    add_callout(doc, "建议文字：", quote)

    add_heading(doc, "七、最终建议", 1)
    add_bullet(doc, "字段保持和干扰捕获率可以保留，但必须补充定义、计算方向和“子实验等权平均”的说明。")
    add_bullet(doc, "Qwen3-8B 不应当被伪装成普通低分样本；它是正常 API 返回但严重输出截断的诊断性异常，应透明说明。")
    add_bullet(doc, "PAC-D 的条件级热图不适合作为核心统计图；保留其机制验证价值，同时弱化其细粒度排名含义。")
    add_bullet(doc, "先处理图例重叠、PAC-D 图形替换和指标定义，再将图表插入正文，会比直接改数更稳妥。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
